"""LocalDocIngestor — Sub-project E: local document ingestion pipeline."""
from __future__ import annotations

import hashlib
import logging
import os
import pathlib
import re
import sqlite3
import threading
import uuid
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import Any

logger = logging.getLogger("rag-bridge.local_doc_ingestor")

DOCS_COLLECTION = "docs_reference"
MAX_FILE_BYTES = 50 * 1024 * 1024  # 50 MB

# Lazy import guards
try:
    import pypdf  # type: ignore[import]
except ImportError:
    pypdf = None  # type: ignore[assignment]

try:
    import docx  # type: ignore[import]  # python-docx
except ImportError:
    docx = None  # type: ignore[assignment]

try:
    import litellm  # type: ignore[import]
except ImportError:
    litellm = None  # type: ignore[assignment]


@dataclass
class IngestResult:
    status: str          # 'indexed' | 'skipped' | 'error' | 'disabled'
    doc_id: str = ""
    file_path: str = ""
    file_type: str = ""
    chunks_count: int = 0
    title: str = ""
    tags: list[str] = field(default_factory=list)
    reason: str = ""
    error_message: str = ""

    def as_dict(self) -> dict:
        return {
            "status": self.status,
            "doc_id": self.doc_id,
            "file_path": self.file_path,
            "file_type": self.file_type,
            "chunks_count": self.chunks_count,
            "title": self.title,
            "tags": self.tags,
            "reason": self.reason,
            "error_message": self.error_message,
        }


class LocalDocIngestor:
    """Pipeline: detect format -> extract text -> chunk -> PII filter -> embed -> Qdrant upsert."""

    def __init__(
        self,
        state_dir: str | pathlib.Path,
        qdrant_client: Any,
    ) -> None:
        self._state_dir = pathlib.Path(state_dir)
        self._qdrant = qdrant_client
        self._db_path = self._state_dir / "scheduler.db"

        self._enabled = os.getenv("LOCAL_DOCS_ENABLED", "false").lower() in ("1", "true", "yes")
        self._watch_path = pathlib.Path(
            os.getenv("LOCAL_DOCS_WATCH_PATH", "/opt/nanobot-stack/watched-docs/")
        ).resolve()
        self._chunk_size = int(os.getenv("LOCAL_DOCS_CHUNK_SIZE", "512"))
        self._chunk_overlap = int(os.getenv("LOCAL_DOCS_CHUNK_OVERLAP", "50"))
        self._formats = set(
            os.getenv("LOCAL_DOCS_FORMATS", "pdf,md,txt,docx").lower().split(",")
        )
        self._watch_path.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------
    # DB helpers
    # ------------------------------------------------------------------

    def _connect(self) -> sqlite3.Connection:
        conn = sqlite3.connect(str(self._db_path))
        conn.row_factory = sqlite3.Row
        return conn

    def _now(self) -> str:
        return datetime.now(timezone.utc).isoformat()

    # ------------------------------------------------------------------
    # Core helpers
    # ------------------------------------------------------------------

    def _hash_file(self, file_path: str) -> str:
        """Return SHA-256 hex digest of file binary content."""
        h = hashlib.sha256()
        with open(file_path, "rb") as fh:
            for chunk in iter(lambda: fh.read(65536), b""):
                h.update(chunk)
        return h.hexdigest()

    def _detect_format(self, file_path: str) -> str | None:
        ext = pathlib.Path(file_path).suffix.lower().lstrip(".")
        return ext if ext in {"pdf", "md", "txt", "docx"} else None

    def _is_already_indexed(self, file_path: str, file_hash: str) -> bool:
        """Return True if file_path exists in log with the same hash (no change)."""
        db = self._connect()
        try:
            row = db.execute(
                "SELECT file_hash FROM docs_ingestion_log WHERE file_path = ?",
                (file_path,),
            ).fetchone()
            if row is None:
                return False
            return row["file_hash"] == file_hash
        finally:
            db.close()

    def _assert_within_watch_path(self, file_path: str) -> None:
        resolved = pathlib.Path(file_path).resolve()
        watch_resolved = pathlib.Path(self._watch_path).resolve()
        if not resolved.is_relative_to(watch_resolved):
            raise PermissionError(
                f"Path '{resolved}' is outside the allowed watch path '{self._watch_path}'"
            )

    # ------------------------------------------------------------------
    # Text extraction
    # ------------------------------------------------------------------

    def _extract_text(self, file_path: str, fmt: str) -> str:
        p = pathlib.Path(file_path)
        if fmt == "txt":
            return p.read_text(encoding="utf-8", errors="replace")
        if fmt == "md":
            return p.read_text(encoding="utf-8")
        if fmt == "pdf":
            if pypdf is None:
                raise ImportError("pypdf is required for PDF ingestion")
            reader = pypdf.PdfReader(file_path)
            pages_text = []
            for page in reader.pages:
                t = page.extract_text() or ""
                pages_text.append(t)
            return " ".join(pages_text)
        if fmt == "docx":
            if docx is None:
                raise ImportError("python-docx is required for DOCX ingestion")
            document = docx.Document(file_path)
            return "\n".join(p_item.text for p_item in document.paragraphs)
        raise ValueError(f"_extract_text called with unsupported fmt={fmt!r}")

    # ------------------------------------------------------------------
    # Metadata extraction
    # ------------------------------------------------------------------

    _FR_STOPWORDS = frozenset(["de", "la", "le", "les", "un", "une", "du", "des"])

    def _extract_metadata(self, file_path: str) -> dict:
        p = pathlib.Path(file_path)
        fmt = self._detect_format(file_path) or "txt"
        title = self._extract_title(p, fmt)
        tags = self._extract_tags(p)
        return {
            "title": title,
            "tags": tags,
            "source_path": str(p.resolve()),
            "file_type": fmt,
        }

    def _extract_title(self, p: pathlib.Path, fmt: str) -> str:
        try:
            if fmt == "md":
                for line in p.read_text(encoding="utf-8").splitlines():
                    if line.startswith("# "):
                        return line[2:].strip()
            elif fmt == "txt":
                for line in p.read_text(encoding="utf-8", errors="replace").splitlines():
                    if line.strip():
                        return line.strip()[:80]
            elif fmt == "pdf":
                if pypdf is None:
                    return p.stem
                reader = pypdf.PdfReader(str(p))
                meta = reader.metadata or {}
                title_val = meta.get("/Title") or meta.get("title") or ""
                if title_val:
                    return str(title_val).strip()
                # Fallback to first non-empty text line
                for page in reader.pages:
                    text = page.extract_text() or ""
                    for line in text.splitlines():
                        if line.strip():
                            return line.strip()[:80]
            elif fmt == "docx":
                if docx is None:
                    return p.stem
                document = docx.Document(str(p))
                core_title = document.core_properties.title or ""
                if core_title.strip():
                    return core_title.strip()
                for para in document.paragraphs:
                    if para.text.strip():
                        return para.text.strip()[:80]
        except Exception:
            pass
        return p.stem

    def _extract_tags(self, p: pathlib.Path) -> list[str]:
        try:
            rel = p.resolve().relative_to(self._watch_path)
        except ValueError:
            rel = p
        parts = list(rel.parts)
        tokens: list[str] = []
        for part in parts:
            tokens.extend(re.split(r"[/\-_ ]", pathlib.Path(part).stem))
        # Add file_type as last tag
        ext = p.suffix.lower().lstrip(".")
        if ext:
            tokens.append(ext)
        seen: set[str] = set()
        result: list[str] = []
        for t in tokens:
            t = t.lower()
            if t and t not in self._FR_STOPWORDS and t not in seen:
                seen.add(t)
                result.append(t)
        return result

    # ------------------------------------------------------------------
    # Chunking
    # ------------------------------------------------------------------

    def _chunk_text(self, text: str, _title: str) -> list[str]:
        """Sliding-window chunker on sentence boundaries."""
        overlap_chars = self._chunk_overlap * 4
        size_chars = self._chunk_size * 4

        # Short text -> single chunk
        if len(text) <= overlap_chars:
            return [text]

        sentences = re.split(r"(?<=\.)\s+|(?<=\.\n)|\n\n", text)
        sentences = [s for s in sentences if s.strip()]

        chunks: list[str] = []
        current_parts: list[str] = []
        current_len = 0

        for sent in sentences:
            if current_len + len(sent) > size_chars and current_parts:
                chunk_text = " ".join(current_parts)
                chunks.append(chunk_text)
                # Overlap: keep last overlap_chars worth of content
                overlap_text = chunk_text[-overlap_chars:]
                current_parts = [overlap_text]
                current_len = len(overlap_text)
            current_parts.append(sent)
            current_len += len(sent)

        if current_parts:
            chunks.append(" ".join(current_parts))

        return chunks if chunks else [text]

    # ------------------------------------------------------------------
    # Embedding
    # ------------------------------------------------------------------

    def _embed(self, text: str) -> list[float] | None:
        try:
            if litellm is None:
                raise ImportError("litellm is required for embedding")
            resp = litellm.embedding(model="text-embedding-3-small", input=[text])
            return resp["data"][0]["embedding"]
        except Exception as exc:
            logger.warning("Embedding failed: %s", exc)
            return None

    # ------------------------------------------------------------------
    # Qdrant operations
    # ------------------------------------------------------------------

    def _embed_and_upsert(self, chunks: list[str], metadata: dict) -> int:
        """Embed chunks with PII filter applied, upsert to Qdrant. Returns count of points upserted."""
        from pii_filter import redact_for_ingest
        try:
            from qdrant_client.models import PointStruct  # type: ignore[import]
        except ImportError as exc:
            logger.error("qdrant_client not available: %s", exc)
            return 0

        points = []
        doc_id = metadata["doc_id"]
        total = len(chunks)
        for i, chunk in enumerate(chunks):
            filtered_chunk, _ = redact_for_ingest(chunk)
            vector = self._embed(filtered_chunk)
            if vector is None:
                continue
            payload = {
                "source_path": metadata["source_path"],
                "file_type": metadata["file_type"],
                "file_hash": metadata["file_hash"],
                "chunk_index": i,
                "total_chunks": total,
                "title": metadata["title"],
                "tags": metadata["tags"],
                "ingested_at": metadata["ingested_at"],
                "doc_id": doc_id,
                "text": filtered_chunk,
            }
            point_id = str(uuid.uuid5(uuid.UUID(doc_id), str(i)))
            points.append(PointStruct(id=point_id, vector=vector, payload=payload))

        upserted = 0
        for batch_start in range(0, len(points), 100):
            batch = points[batch_start: batch_start + 100]
            self._qdrant.upsert(collection_name=DOCS_COLLECTION, points=batch)
            upserted += len(batch)

        return upserted

    def _delete_qdrant_chunks(self, doc_id: str) -> None:
        try:
            from qdrant_client.models import Filter, FieldCondition, MatchValue, FilterSelector  # type: ignore[import]
            self._qdrant.delete(
                collection_name=DOCS_COLLECTION,
                points_selector=FilterSelector(
                    filter=Filter(must=[FieldCondition(key="doc_id", match=MatchValue(value=doc_id))])
                ),
            )
        except Exception as exc:
            logger.warning("Qdrant delete failed for doc_id=%s: %s", doc_id, exc)

    # ------------------------------------------------------------------
    # DB log operations
    # ------------------------------------------------------------------

    def _get_existing_doc_id(self, file_path: str) -> str | None:
        db = self._connect()
        try:
            row = db.execute(
                "SELECT id FROM docs_ingestion_log WHERE file_path = ?", (file_path,)
            ).fetchone()
            return row["id"] if row else None
        finally:
            db.close()

    def _upsert_log(
        self, doc_id: str, file_path: str, file_hash: str,
        file_type: str, chunks_count: int, status: str, error_message: str = ""
    ) -> None:
        now = self._now()
        db = self._connect()
        try:
            existing = db.execute(
                "SELECT created_at FROM docs_ingestion_log WHERE id = ? OR file_path = ?",
                (doc_id, file_path),
            ).fetchone()
            created_at = existing["created_at"] if existing else now
            db.execute(
                """
                INSERT OR REPLACE INTO docs_ingestion_log
                  (id, file_path, file_hash, file_type, chunks_count, status, error_message, last_indexed, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (doc_id, file_path, file_hash, file_type, chunks_count,
                 status, error_message or None, now, created_at),
            )
            db.commit()
        finally:
            db.close()

    def _log_error(self, file_path: str, message: str) -> None:
        doc_id = self._get_existing_doc_id(file_path) or str(uuid.uuid4())
        fmt = self._detect_format(file_path) or "unknown"
        self._upsert_log(doc_id, file_path, "", fmt, 0, "error", message)

    def _update_log_skipped(self, file_path: str) -> None:
        now = self._now()
        db = self._connect()
        try:
            db.execute(
                "UPDATE docs_ingestion_log SET status='skipped', last_indexed=? WHERE file_path=?",
                (now, file_path),
            )
            db.commit()
        finally:
            db.close()

    # ------------------------------------------------------------------
    # Public entrypoint
    # ------------------------------------------------------------------

    def ingest_file(self, file_path: str) -> dict:
        if not self._enabled:
            return IngestResult(status="disabled").as_dict()
        self._assert_within_watch_path(file_path)

        # File size guard
        size = pathlib.Path(file_path).stat().st_size
        if size > MAX_FILE_BYTES:
            msg = f"File too large: {size} bytes (max {MAX_FILE_BYTES})"
            self._log_error(file_path, msg)
            return IngestResult(status="error", file_path=file_path, error_message=msg).as_dict()

        fmt = self._detect_format(file_path)
        if fmt is None or fmt not in self._formats:
            return IngestResult(status="skipped", file_path=file_path, reason="unsupported_format").as_dict()

        file_hash = self._hash_file(file_path)
        if self._is_already_indexed(file_path, file_hash):
            self._update_log_skipped(file_path)
            return IngestResult(status="skipped", file_path=file_path, reason="same_hash").as_dict()

        # Check for modified file (different hash but known path)
        existing_doc_id = self._get_existing_doc_id(file_path)
        if existing_doc_id:
            self._delete_qdrant_chunks(existing_doc_id)

        try:
            text = self._extract_text(file_path, fmt)
        except Exception as exc:
            msg = f"Extraction error: {exc}"
            logger.exception("extract_text failed for %s", file_path)
            self._log_error(file_path, msg)
            return IngestResult(status="error", file_path=file_path, error_message=msg).as_dict()

        meta = self._extract_metadata(file_path)
        chunks = self._chunk_text(text, meta["title"])

        doc_id = existing_doc_id or str(uuid.uuid4())
        metadata = {
            "doc_id": doc_id,
            "source_path": str(pathlib.Path(file_path).resolve()),
            "file_type": fmt,
            "file_hash": file_hash,
            "title": meta["title"],
            "tags": meta["tags"],
            "ingested_at": self._now(),
        }
        self._embed_and_upsert(chunks, metadata)

        self._upsert_log(doc_id, file_path, file_hash, fmt, len(chunks), "indexed")
        return IngestResult(
            status="indexed",
            doc_id=doc_id,
            file_path=file_path,
            file_type=fmt,
            chunks_count=len(chunks),
            title=meta["title"],
            tags=meta["tags"],
        ).as_dict()

    def delete_document(self, doc_id: str) -> bool:
        db = self._connect()
        try:
            row = db.execute(
                "SELECT status FROM docs_ingestion_log WHERE id=?", (doc_id,)
            ).fetchone()
            if row is None or row["status"] == "deleted":
                return False
        finally:
            db.close()

        self._delete_qdrant_chunks(doc_id)
        now = self._now()
        db = self._connect()
        try:
            db.execute(
                "UPDATE docs_ingestion_log SET status='deleted', last_indexed=? WHERE id=?",
                (now, doc_id),
            )
            db.commit()
        finally:
            db.close()
        return True

    def list_documents(self, limit: int = 100, offset: int = 0,
                       file_type: str | None = None, status: str | None = None) -> list[dict]:
        db = self._connect()
        try:
            query = "SELECT * FROM docs_ingestion_log WHERE status != 'deleted'"
            params: list[Any] = []
            if file_type:
                query += " AND file_type=?"
                params.append(file_type)
            if status:
                query += " AND status=?"
                params.append(status)
            query += " ORDER BY last_indexed DESC LIMIT ? OFFSET ?"
            params += [limit, offset]
            rows = db.execute(query, params).fetchall()
            return [dict(r) for r in rows]
        finally:
            db.close()

    def get_status(self) -> dict:
        db = self._connect()
        try:
            total_row = db.execute(
                "SELECT COUNT(*) FROM docs_ingestion_log WHERE status='indexed'"
            ).fetchone()
            total_files = total_row[0] if total_row else 0

            chunks_row = db.execute(
                "SELECT SUM(chunks_count) FROM docs_ingestion_log WHERE status='indexed'"
            ).fetchone()
            total_chunks = chunks_row[0] or 0

            last_row = db.execute(
                "SELECT MAX(last_indexed) FROM docs_ingestion_log WHERE status='indexed'"
            ).fetchone()
            last_indexed = last_row[0] if last_row else None

            type_rows = db.execute(
                "SELECT file_type, COUNT(*) as files, SUM(chunks_count) as chunks "
                "FROM docs_ingestion_log WHERE status='indexed' GROUP BY file_type"
            ).fetchall()
            breakdown = {
                r["file_type"]: {"files": r["files"], "chunks": r["chunks"] or 0}
                for r in type_rows
            }
        finally:
            db.close()

        return {
            "enabled": self._enabled,
            "watch_path": str(self._watch_path),
            "total_files": total_files,
            "total_chunks": total_chunks,
            "last_indexed": last_indexed,
            "breakdown": breakdown,
            "watcher_running": False,  # updated by LocalDocWatcher after start
        }

    def ingest_directory(self, path: str | None = None) -> list[dict]:
        if not self._enabled:
            return [{"status": "disabled"}]
        scan_path = pathlib.Path(path) if path else self._watch_path
        results = []
        for fp in scan_path.rglob("*"):
            if not fp.is_file():
                continue
            fmt = self._detect_format(str(fp))
            if fmt is None or fmt not in self._formats:
                continue
            try:
                r = self.ingest_file(str(fp))
                results.append(r)
            except Exception as exc:
                logger.warning("ingest_directory: error on %s: %s", fp, exc)
                results.append({"status": "error", "file_path": str(fp), "error_message": str(exc)})
        return results


class LocalDocWatcher:
    """Watchdog-based file system monitor. Calls LocalDocIngestor on file events."""

    # Temp/hidden file patterns to ignore
    _IGNORE_EXTENSIONS = frozenset([".tmp", ".part", ".swp"])

    def __init__(self) -> None:
        self._timers: dict[str, Any] = {}
        self._lock = threading.Lock()
        self._observer: Any = None
        self._running = False

    def _is_temp_file(self, path: str) -> bool:
        p = pathlib.Path(path)
        if p.name.startswith("."):
            return True
        if p.suffix.lower() in self._IGNORE_EXTENSIONS:
            return True
        return False

    def _handle_event(self, event: Any, ingestor: LocalDocIngestor) -> None:
        if event.is_directory:
            return
        src_path = event.src_path
        if self._is_temp_file(src_path):
            return

        def _fire():
            with self._lock:
                self._timers.pop(src_path, None)
            try:
                ingestor.ingest_file(src_path)
            except Exception as exc:
                logger.warning("Watcher ingest_file failed for %s: %s", src_path, exc)

        with self._lock:
            existing = self._timers.get(src_path)
            if existing is not None:
                existing.cancel()
            t = threading.Timer(2.0, _fire)
            self._timers[src_path] = t
            t.start()

    def start(self, path: str, ingestor: LocalDocIngestor) -> None:
        try:
            from watchdog.observers import Observer  # type: ignore[import]
            from watchdog.events import FileSystemEventHandler  # type: ignore[import]
        except ImportError:
            logger.warning("watchdog not installed — LocalDocWatcher not started")
            return

        watcher_ref = self

        class _Handler(FileSystemEventHandler):
            def on_created(self, event):
                """Handle file creation events."""
                watcher_ref._handle_event(event, ingestor)  # pylint: disable=protected-access

            def on_modified(self, event):
                """Handle file modification events."""
                watcher_ref._handle_event(event, ingestor)  # pylint: disable=protected-access

        self._observer = Observer()
        self._observer.schedule(_Handler(), path, recursive=True)
        self._observer.daemon = True
        self._observer.start()
        self._running = True
        logger.info("LocalDocWatcher started on %s", path)

    def stop(self) -> None:
        """Stop the watchdog observer."""
        if self._observer and self._running:
            self._observer.stop()
            self._observer.join()
            self._running = False
            logger.info("LocalDocWatcher stopped")

    @property
    def is_running(self) -> bool:
        """Return True if the observer is running."""
        return self._running
